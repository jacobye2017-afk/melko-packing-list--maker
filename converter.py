"""
Packing List Maker - Core converter

Mode 1: 统一 packing list 格式（派送方案单, 19-column layout）

Target layout (columns):
  A  NO
  B  SHIP ID（分货号）
  C  METHOD\n渠道
  D  CTNS件数
  E  CBM
  F  WEIGHR（KGS）
  G  CTN/LBS           (computed: kg × 2.2 / ctns)
  H  总LBS              (computed: kg × 2.2)
  I  DESCRIPTION（品名）
  J  FBA CODE（仓库代码）
  K  ADDRESS（地址）
  L  FBA ID
  M  REFERENCE ID
  N  装柜顺序
  O  预计到达日期
  P  TRUCKING NO.
  Q  RATE
  R  备注
  S+ unmapped source columns dumped verbatim

Supports:
  - .xlsx (openpyxl) and .xls (xlrd) source files
  - Multiple customer source-header conventions via alias matching
  - Auto-detection of header row (skips container/metadata rows at top)
"""

import os
import re
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet

try:
    import xlrd  # for .xls
    _HAS_XLRD = True
except ImportError:
    _HAS_XLRD = False


# =============================================================================
# Constants
# =============================================================================

KG_TO_LBS = 2.2

METHOD_FILLS = {
    "FEDEX":       "7030A0",  # purple
    "UPS":         "806000",  # dark amber/brown
    "HOLD":        "FF0000",  # red
    "PICK UP":     "FF0000",
    "PICKUP":      "FF0000",
    "RELABEL":     "FF0000",
    "LOCAL":       "FFFF00",  # yellow
    "LTL":         "FFFF00",
    "FF":          "FFFF00",
    "FULFILLMENT": "FFFF00",
    "JOYCE":       "FFFF00",
    "JOCYE":       "FFFF00",
}

METHODS_WHITE_FONT = {"FEDEX", "UPS"}

HEADER_FILL = "D9E2F3"
FONT_NAME = "微软雅黑"

# 19-column widths
COLUMN_WIDTHS = {
    "A": 6.0,   "B": 33.0,  "C": 16.5,  "D": 10.66,
    "E": 10.0,  "F": 12.0,  "G": 12.0,  "H": 12.0,
    "I": 22.0,  "J": 17.5,  "K": 29.5,  "L": 17.33,
    "M": 18.0,  "N": 12.0,  "O": 14.0,  "P": 16.0,
    "Q": 10.0,  "R": 22.0,
}

# Methods that should NOT merge FBA CODE cells (keep each row independent)
NO_MERGE_METHODS = {"HOLD", "RELABEL", "PICK UP", "PICKUP"}
EXTRAS_DEFAULT_WIDTH = 14.0

ROW_H_TITLE = 30.0
ROW_H_SUB   = 28.0
ROW_H_HEAD  = 42.0
ROW_H_DATA  = 30.0
ROW_H_ADDR  = 40.0

# Target headers (row 3)
HEADERS = [
    "NO",
    "SHIP ID（分货号）",
    "METHOD\n渠道",
    "CTNS件数",
    "CBM",
    "WEIGHR（KGS）",
    "CTN/LBS",
    "总LBS",
    "DESCRIPTION（品名）",
    "FBA CODE（仓库代码）",
    "ADDRESS（地址）",
    "FBA ID",
    "REFERENCE ID",
    "装柜顺序",
    "预计到达日期",
    "TRUCKING NO.",
    "RATE",
    "备注",
]
N_MAIN_COLS = len(HEADERS)  # 18
EXTRAS_START_COL = N_MAIN_COLS + 1  # S = 19


# ---- Source-header alias table --------------------------------------------
HEADER_ALIASES = {
    "ship_id":     [
        "唛头", "ship id", "ship id(分货号)", "分货号", "shipid", "ship id（分货号）",
        "ship id（唛头）",
    ],
    "method":      ["派送方式", "method", "ship method"],
    "channel":     ["渠道", "channel"],
    "ctns":        ["总件数", "件数", "件数ctns", "ctns件数", "ctns"],
    "cbm":         ["总cbm", "cbm", "体积cbm"],
    "kg":          ["总kg", "总重", "kg", "weight", "weighr(kgs)", "weighr（kgs）", "weighr"],
    "lbs":         ["总lbs", "lbs"],
    "description": ["中文品名", "品名", "description", "description(品名)", "description（品名）"],
    "fba_code":    [
        "fbx 仓库代码(私人地址填priv)", "fbx仓库代码", "fba code", "仓库代码",
        "fba code(仓库代码)", "fba code（仓库代码）",
    ],
    "address":     ["地址", "address", "address(地址)", "address（地址）"],
    "fba_id":      ["fbx货件编号", "fba id", "fba货件编号", "货件编号"],
    "refid_alt":   ["refid", "ref id", "ref-id", "reference id"],
    "remark":      ["备注", "remark", "note", "notes"],
    "ref_primary": ["快递主单号", "快递单号", "主单号", "tracking", "tracking no", "tracking no."],
    # KINGONE target fields
    "loading_order": ["装柜顺序", "window time", "打板数"],
    "eta_date":    ["预计到达日期", "eta"],
    "trucking_no": ["trucking no.", "trucking no", "trucking"],
    "rate":        ["rate", "费率"],
}

# Canonicals whose source values go into main target columns (so NOT copied to extras)
MAPPED_CANONICALS = set(HEADER_ALIASES.keys())

# Source headers to always ignore (not copied to extras, not used for data).
# These are noise columns that appear in some customer templates.
SKIP_HEADER_KEYS = {
    "no",           # source serial-number column (we re-number in target)
    "单件pcs",      # per-unit count, not used
}


# =============================================================================
# Helpers
# =============================================================================

def _norm_str(v):
    if v is None:
        return ""
    if isinstance(v, float) and v != v:  # NaN
        return ""
    s = str(v).strip()
    return s.replace("\xa0", " ")  # normalize NBSP


def _norm_header_key(name):
    if not name:
        return ""
    s = str(name).strip().lower()
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("\xa0", " ")
    s = re.sub(r"\s+", " ", s)
    return s


def _norm_method(m):
    """Normalize method string to canonical form (upper case, standardized)."""
    s = _norm_str(m).upper()
    if not s:
        return ""
    # Normalize common variants
    if "一件代发" in s or "一件代发" in _norm_str(m):
        return "FF"
    if s in ("FEDEX", "FEDEX "):
        return "FEDEX"
    return s


def _infer_method(row, header_map):
    """Derive the delivery method for a row.
    Priority: explicit 派送方式 → 渠道 → empty."""
    method_key = header_map.get("method")
    channel_key = header_map.get("channel")
    m = _norm_str(row.get(method_key)) if method_key else ""
    if not m:
        m = _norm_str(row.get(channel_key)) if channel_key else ""
    return _norm_method(m) or m.upper()


def _method_fill(method):
    return METHOD_FILLS.get(_norm_method(method))


def _is_white_font(method):
    return _norm_method(method) in METHODS_WHITE_FONT


def _thin_border():
    s = Side(style="thin", color="000000")
    return Border(left=s, right=s, top=s, bottom=s)


def _to_num(v, default=0):
    if v is None or v == "":
        return default
    if isinstance(v, (int, float)):
        return v
    try:
        return float(str(v).strip())
    except (TypeError, ValueError):
        return default


def _extract_container(filename, fallback_text=""):
    """Extract container number from filename or text.
    4 uppercase letters + 7 digits, case-insensitive."""
    for text in (os.path.basename(filename), fallback_text):
        if not text:
            continue
        matches = re.findall(r"(?<![A-Za-z0-9])[A-Za-z]{4}\d{7}(?![0-9])", text)
        if matches:
            return matches[0].upper()
    stem = os.path.splitext(os.path.basename(filename))[0]
    return re.split(r"[\s\-_+]", stem)[0].upper()


# =============================================================================
# Read source (.xlsx or .xls)
# =============================================================================

def _load_sheet_as_grid(src_path):
    """Return a 2D list of values regardless of .xls or .xlsx format.
    Also returns the sheet/workbook name-of-first-sheet.
    """
    ext = os.path.splitext(src_path)[1].lower()
    if ext == ".xlsx":
        wb = openpyxl.load_workbook(src_path, data_only=True)
        # Prefer 导入数据 / 派送表 sheets if present, else first non-hidden
        preferred = ["导入数据", "派送表", "总表", "packing list"]
        ws = None
        for name in preferred:
            if name in wb.sheetnames:
                ws = wb[name]
                break
        if ws is None:
            ws = wb.worksheets[0]
        grid = []
        for row in ws.iter_rows(values_only=True):
            grid.append(list(row))
        return grid, ws.title
    elif ext == ".xls":
        if not _HAS_XLRD:
            raise RuntimeError("Reading .xls requires the 'xlrd' package. Run: pip install xlrd")
        wb = xlrd.open_workbook(src_path)
        preferred = ["导入数据", "派送表", "总表"]
        ws = None
        for name in preferred:
            if name in wb.sheet_names():
                ws = wb.sheet_by_name(name)
                break
        if ws is None:
            ws = wb.sheet_by_index(0)
        grid = []
        for r in range(ws.nrows):
            row = []
            for c in range(ws.ncols):
                v = ws.cell_value(r, c)
                ct = ws.cell_type(r, c)
                if ct == xlrd.XL_CELL_EMPTY:
                    row.append(None)
                elif ct == xlrd.XL_CELL_DATE:
                    from datetime import datetime
                    try:
                        row.append(datetime(*xlrd.xldate_as_tuple(v, wb.datemode)))
                    except Exception:
                        row.append(v)
                elif ct == xlrd.XL_CELL_NUMBER:
                    # Keep integers as ints when possible
                    row.append(int(v) if v == int(v) else v)
                else:
                    row.append(v if v != "" else None)
            grid.append(row)
        return grid, ws.name
    else:
        raise RuntimeError(f"Unsupported file extension: {ext}")


def _detect_header_row(grid, max_scan=8):
    """Scan the first max_scan rows for the one that looks like a header row.
    Heuristic: row with the highest count of recognized header keywords."""
    recognized = set()
    for canonical, aliases in HEADER_ALIASES.items():
        for a in aliases:
            recognized.add(_norm_header_key(a))

    best_idx = 0
    best_score = -1
    for i, row in enumerate(grid[:max_scan]):
        score = 0
        for cell in row:
            k = _norm_header_key(cell)
            if k in recognized:
                score += 1
        if score > best_score and score >= 3:
            best_score = score
            best_idx = i
    return best_idx if best_score >= 3 else 0


def _extract_metadata(grid, header_row_idx):
    """Extract container, ETD, ETA, etc from rows above the header row."""
    meta = {"container": "", "etd": "", "eta": ""}
    for row in grid[:header_row_idx]:
        for i, cell in enumerate(row):
            if cell is None:
                continue
            s = _norm_str(cell)
            if not s:
                continue
            # container?
            cm = re.search(r"(?<![A-Za-z0-9])[A-Za-z]{4}\d{7}(?![0-9])", s)
            if cm and not meta["container"]:
                meta["container"] = cm.group(0).upper()
            # ETD label near a date
            if s.upper().startswith("ETD") and i + 1 < len(row):
                val = row[i + 1]
                if val:
                    meta["etd"] = _norm_str(val)
            if s.upper().startswith("ETA") and i + 1 < len(row):
                val = row[i + 1]
                if val:
                    meta["eta"] = _norm_str(val)
    return meta


def read_source(src_path):
    """Read a source spreadsheet (.xlsx or .xls) and return:
      raw_headers, rows (list of dicts), header_map (canonical → actual header),
      header_to_canon (actual header → canonical or None), metadata (dict).
    """
    grid, _ = _load_sheet_as_grid(src_path)
    if not grid:
        return [], [], {}, {}, {}

    header_idx = _detect_header_row(grid)
    metadata = _extract_metadata(grid, header_idx)

    raw_headers = [(_norm_str(c) or f"col{i}") for i, c in enumerate(grid[header_idx])]

    # Build alias lookup
    alias_lookup = {}
    for canonical, aliases in HEADER_ALIASES.items():
        for a in aliases:
            alias_lookup[_norm_header_key(a)] = canonical

    header_map = {}
    header_to_canon = {}
    for h in raw_headers:
        key = _norm_header_key(h)
        canon = alias_lookup.get(key)
        header_to_canon[h] = canon
        if canon and canon not in header_map:
            header_map[canon] = h

    # Read data rows (below header_idx)
    rows = []
    ship_id_key = header_map.get("ship_id")
    for raw_row in grid[header_idx + 1:]:
        if not any(c not in (None, "") for c in raw_row):
            continue
        d = {}
        for h, v in zip(raw_headers, raw_row):
            d[h] = v
        mark = _norm_str(d.get(ship_id_key)) if ship_id_key else ""
        if not mark:
            # Also skip rows that look like a TOTAL row
            row_vals = [_norm_str(x).upper() for x in raw_row if x not in (None, "")]
            if any("总计" in v or "TOTAL" in v for v in row_vals):
                continue
            # Skip truly empty
            continue
        if _norm_str(mark).upper() in ("TOTAL", "总计"):
            continue
        rows.append(d)

    return raw_headers, rows, header_map, header_to_canon, metadata


# =============================================================================
# Transform / sort
# =============================================================================

def _group_priority(method):
    m = _norm_method(method)
    if m == "FEDEX":
        return 1
    if m == "UPS":
        return 2
    if m in ("PICK UP", "PICKUP"):
        return 3
    if m in ("HOLD", "RELABEL"):
        return 4
    if m == "TRUCK":
        return 5
    return 6


def sort_rows(rows, header_map):
    fba_key = header_map.get("fba_code")
    channel_key = header_map.get("channel")
    method_key = header_map.get("method")

    def method_of(r):
        m = _norm_str(r.get(method_key)) if method_key else ""
        if not m and channel_key:
            m = _norm_str(r.get(channel_key))
        return _norm_method(m)

    # First-seen warehouse order for TRUCK rows
    truck_wh_order = {}
    for r in rows:
        if method_of(r) == "TRUCK":
            wh = _norm_str(r.get(fba_key)) if fba_key else ""
            if wh and wh not in truck_wh_order:
                truck_wh_order[wh] = len(truck_wh_order)

    # First-seen "others" method order
    other_order = {}
    for r in rows:
        if _group_priority(method_of(r)) == 6:
            m = method_of(r)
            if m and m not in other_order:
                other_order[m] = len(other_order)

    def key(idx_row):
        idx, r = idx_row
        m = method_of(r)
        gp = _group_priority(m)
        if gp == 5:
            wh = _norm_str(r.get(fba_key)) if fba_key else ""
            return (gp, truck_wh_order.get(wh, 9999), idx)
        if gp == 6:
            return (gp, other_order.get(m, 9999), idx)
        return (gp, 0, idx)

    indexed = list(enumerate(rows))
    indexed.sort(key=key)
    return [r for _, r in indexed]


def transform_row(src, header_map):
    def g(canon, default=None):
        key = header_map.get(canon)
        if key is None:
            return default
        return src.get(key, default)

    ctns = _to_num(g("ctns"))
    kg = _to_num(g("kg"))
    cbm = _to_num(g("cbm"))
    lbs = round(kg * KG_TO_LBS, 3)
    ctn_lbs = round(lbs / ctns, 6) if ctns else 0

    ref_primary = _norm_str(g("ref_primary"))
    ref_alt = _norm_str(g("refid_alt"))
    ref = ref_primary or ref_alt
    if ref == "/":
        ref = ""

    remark = _norm_str(g("remark"))
    address = _norm_str(g("address"))
    if remark and remark == address:
        remark = ""

    method = _infer_method(src, header_map)

    # Normalize FBA CODE: anything containing "一件代发" → "fulfillment"
    fba_code_raw = _norm_str(g("fba_code"))
    if "一件代发" in fba_code_raw:
        fba_code_raw = "fulfillment"

    return {
        "no": "",
        "ship_id": _norm_str(g("ship_id")),
        "method": method,
        "ctns": int(ctns) if ctns == int(ctns) else ctns,
        "cbm": cbm,
        "kg": kg,
        "ctn_lbs": ctn_lbs,
        "total_lbs": lbs,
        "description": _norm_str(g("description")),
        "fba_code": fba_code_raw,
        "address": address,
        "fba_id": _norm_str(g("fba_id")),
        "reference_id": ref,
        "loading_order": _norm_str(g("loading_order")),
        "eta_date": _norm_str(g("eta_date")),
        "trucking_no": _norm_str(g("trucking_no")),
        "rate": _norm_str(g("rate")),
        "remark": remark,
    }


# =============================================================================
# Write Mode 1 output (19-col layout)
# =============================================================================

def _write_header_block(ws: Worksheet, container_no: str, metadata: dict, extras_headers):
    thin = _thin_border()

    # --- Row 1: container + optional 卸柜日期 ---
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=N_MAIN_COLS)
    c = ws.cell(row=1, column=1, value=container_no)
    c.font = Font(name=FONT_NAME, size=28, bold=True, color="000000")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = ROW_H_TITLE

    # --- Row 2: Container NO / ETD / ETA labels ---
    ws.merge_cells("A2:B2")
    ws.cell(row=2, column=1, value="Container NO(集装箱号）")
    ws.merge_cells("C2:D2")
    ws.cell(row=2, column=3, value=container_no)
    ws.cell(row=2, column=5, value="ETD")
    ws.merge_cells("F2:I2")
    ws.cell(row=2, column=6, value=metadata.get("etd", ""))
    ws.cell(row=2, column=10, value="ETA")
    ws.merge_cells(start_row=2, start_column=11, end_row=2, end_column=N_MAIN_COLS)
    ws.cell(row=2, column=11, value=metadata.get("eta", ""))
    ws.row_dimensions[2].height = ROW_H_SUB
    for col in range(1, N_MAIN_COLS + 1):
        cc = ws.cell(row=2, column=col)
        cc.font = Font(name=FONT_NAME, size=12, bold=False, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.border = thin

    # --- Row 3: column headers ---
    for col, h in enumerate(HEADERS, start=1):
        cc = ws.cell(row=3, column=col, value=h)
        cc.font = Font(name=FONT_NAME, size=11, bold=True, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cc.border = thin

    # Extras headers
    for i, h in enumerate(extras_headers):
        col = EXTRAS_START_COL + i
        cc = ws.cell(row=3, column=col, value=h)
        cc.font = Font(name=FONT_NAME, size=11, bold=True, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cc.border = thin
        ws.column_dimensions[get_column_letter(col)].width = EXTRAS_DEFAULT_WIDTH

    ws.row_dimensions[3].height = ROW_H_HEAD


def _apply_data_row_style(ws, row, method, n_extras, is_long=False):
    thin = _thin_border()
    fill_color = _method_fill(method)
    white = _is_white_font(method)

    font = Font(name=FONT_NAME, size=11, bold=True,
                color=("FFFFFF" if white else "000000"))
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    fill = PatternFill("solid", fgColor=fill_color) if fill_color else None

    for col in range(1, N_MAIN_COLS + 1):
        cc = ws.cell(row=row, column=col)
        cc.font = font
        cc.alignment = align
        cc.border = thin
        if fill:
            cc.fill = fill

    extras_font = Font(name=FONT_NAME, size=10, bold=False, color="000000")
    for i in range(n_extras):
        col = EXTRAS_START_COL + i
        cc = ws.cell(row=row, column=col)
        cc.font = extras_font
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.border = thin

    ws.row_dimensions[row].height = ROW_H_ADDR if is_long else ROW_H_DATA


def _write_data_rows(ws: Worksheet, transformed, sorted_source, extras_headers):
    start_row = 4  # data starts at row 4 (row 1 title, 2 labels, 3 headers)
    n_extras = len(extras_headers)

    for i, (r, src) in enumerate(zip(transformed, sorted_source)):
        row = start_row + i
        ws.cell(row=row, column=1, value=r["no"])
        ws.cell(row=row, column=2, value=r["ship_id"])
        ws.cell(row=row, column=3, value=r["method"])
        ws.cell(row=row, column=4, value=r["ctns"])
        ws.cell(row=row, column=5, value=r["cbm"])
        ws.cell(row=row, column=6, value=r["kg"])
        ws.cell(row=row, column=7, value=r["ctn_lbs"])
        ws.cell(row=row, column=8, value=r["total_lbs"])
        ws.cell(row=row, column=9, value=r["description"])
        ws.cell(row=row, column=10, value=r["fba_code"])
        ws.cell(row=row, column=11, value=r["address"])
        ws.cell(row=row, column=12, value=r["fba_id"])
        ws.cell(row=row, column=13, value=r["reference_id"])
        ws.cell(row=row, column=14, value=r["loading_order"])
        ws.cell(row=row, column=15, value=r["eta_date"])
        ws.cell(row=row, column=16, value=r["trucking_no"])
        ws.cell(row=row, column=17, value=r["rate"])
        ws.cell(row=row, column=18, value=r["remark"])

        for j, h in enumerate(extras_headers):
            v = src.get(h)
            ws.cell(row=row, column=EXTRAS_START_COL + j,
                    value=_norm_str(v) if v is not None else "")

        long = False
        if r["address"] and ("\n" in r["address"] or len(r["address"]) > 40):
            long = True
        if r["remark"] and ("\n" in r["remark"] or len(r["remark"]) > 40):
            long = True
        _apply_data_row_style(ws, row, r["method"], n_extras, is_long=long)

    # Number formats
    for i in range(len(transformed)):
        row = start_row + i
        ws.cell(row=row, column=4).number_format = "0"
        ws.cell(row=row, column=5).number_format = "0.0000"
        ws.cell(row=row, column=6).number_format = "0.00"
        ws.cell(row=row, column=7).number_format = "0.000"
        ws.cell(row=row, column=8).number_format = "0.000"

    _merge_consecutive_same(ws, transformed, start_row, col=10, key="fba_code")

    return start_row + len(transformed)


def _merge_consecutive_same(ws, transformed, start_row, col, key):
    """Merge consecutive rows in `col` that have the same value in `key`.
    Skips merging for HOLD/RELABEL/PICK UP rows (each keeps independent FBA CODE)."""
    i = 0
    n = len(transformed)
    while i < n:
        val = transformed[i][key]
        method = transformed[i].get("method", "")
        # Skip rows whose method is in NO_MERGE_METHODS
        if not val or method in NO_MERGE_METHODS:
            i += 1
            continue
        j = i + 1
        while (j < n
               and transformed[j][key] == val
               and transformed[j].get("method", "") not in NO_MERGE_METHODS):
            j += 1
        if j - i >= 2:
            for k in range(i + 1, j):
                ws.cell(row=start_row + k, column=col).value = None
            ws.merge_cells(
                start_row=start_row + i, start_column=col,
                end_row=start_row + j - 1, end_column=col,
            )
        i = j


def _write_total_row(ws: Worksheet, total_row, transformed, n_extras):
    thin = _thin_border()
    sum_ctns = sum(_to_num(r["ctns"]) for r in transformed)
    sum_cbm = round(sum(_to_num(r["cbm"]) for r in transformed), 4)
    sum_kg = round(sum(_to_num(r["kg"]) for r in transformed), 3)
    sum_ctn_lbs = round(sum(_to_num(r["ctn_lbs"]) for r in transformed), 6)
    sum_lbs = round(sum(_to_num(r["total_lbs"]) for r in transformed), 3)

    ws.cell(row=total_row, column=3, value="TOTAL")
    ws.cell(row=total_row, column=4,
            value=int(sum_ctns) if sum_ctns == int(sum_ctns) else sum_ctns)
    ws.cell(row=total_row, column=5, value=sum_cbm)
    ws.cell(row=total_row, column=6, value=sum_kg)
    ws.cell(row=total_row, column=7, value=sum_ctn_lbs)
    ws.cell(row=total_row, column=8, value=sum_lbs)

    red = Font(name=FONT_NAME, size=11, bold=True, color="FF0000")
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for col in range(1, N_MAIN_COLS + 1):
        cc = ws.cell(row=total_row, column=col)
        cc.font = red
        cc.alignment = align
        cc.border = thin
    for i in range(n_extras):
        cc = ws.cell(row=total_row, column=EXTRAS_START_COL + i)
        cc.font = red
        cc.alignment = align
        cc.border = thin
    ws.row_dimensions[total_row].height = ROW_H_DATA

    ws.cell(row=total_row, column=4).number_format = "0"
    ws.cell(row=total_row, column=5).number_format = "0.0000"
    ws.cell(row=total_row, column=6).number_format = "0.00"
    ws.cell(row=total_row, column=7).number_format = "0.000"
    ws.cell(row=total_row, column=8).number_format = "0.000"


def build_mode1(src_path: str, out_path: str, container_no: str = None):
    raw_headers, rows, header_map, header_to_canon, metadata = read_source(src_path)

    if container_no is None:
        container_no = metadata.get("container") or _extract_container(src_path)

    sorted_rows = sort_rows(rows, header_map)
    transformed = [transform_row(r, header_map) for r in sorted_rows]

    # Decide extras: any source header not mapped, filter out fully-empty cols
    # and skip noise columns.
    all_candidates = [
        h for h in raw_headers
        if header_to_canon.get(h) is None
        and h
        and _norm_header_key(h) not in SKIP_HEADER_KEYS
    ]
    extras_headers = []
    for h in all_candidates:
        if any(sr.get(h) not in (None, "") for sr in sorted_rows):
            extras_headers.append(h)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "总表"

    for col_letter, w in COLUMN_WIDTHS.items():
        ws.column_dimensions[col_letter].width = w

    _write_header_block(ws, container_no, metadata, extras_headers)
    last_data = _write_data_rows(ws, transformed, sorted_rows, extras_headers)
    _write_total_row(ws, last_data, transformed, len(extras_headers))

    ws.freeze_panes = "A4"

    wb.save(out_path)
    return transformed, container_no, metadata


# =============================================================================
# BOL generators
# =============================================================================

# --- BOL columns: A-M only (13 cols) ---
BOL_HEADERS = HEADERS[:13]  # NO .. REFERENCE ID
BOL_N_COLS = 13
BOL_COL_WIDTHS = {
    "A": 6.0,   "B": 32.33, "C": 16.5,  "D": 10.66,
    "E": 11.33, "F": 12.5,  "G": 12.0,  "H": 15.5,
    "I": 20.83, "J": 17.5,  "K": 15.66, "L": 28.5, "M": 28.83,
}
BOL_KEY_ORDER = [
    "no", "ship_id", "method", "ctns", "cbm", "kg",
    "ctn_lbs", "total_lbs", "description", "fba_code",
    "address", "fba_id", "reference_id",
]


def _write_bol_excel(transformed, container_no, title_text, out_path):
    """Shared writer for BOL packing list and hold list (A-M, 13 cols)."""
    thin = _thin_border()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "总表"

    for col_letter, w in BOL_COL_WIDTHS.items():
        ws.column_dimensions[col_letter].width = w

    # Row 1: title
    ws.merge_cells("A1:M1")
    c = ws.cell(row=1, column=1, value=title_text)
    c.font = Font(name=FONT_NAME, size=36, bold=True, color="000000")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 60

    # Row 2: Container / ETD / ETA
    ws.merge_cells("A2:B2")
    ws.cell(row=2, column=1, value="Container NO(集装箱号）")
    ws.merge_cells("C2:D2")
    ws.cell(row=2, column=3, value=container_no)
    ws.cell(row=2, column=5, value="ETD")
    ws.merge_cells("F2:I2")
    ws.cell(row=2, column=10, value="ETA")
    ws.merge_cells("K2:M2")
    ws.row_dimensions[2].height = 39
    for col in range(1, 14):
        cc = ws.cell(row=2, column=col)
        cc.font = Font(name=FONT_NAME, size=16, bold=False, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.border = thin

    # Row 3: 提柜日期 / 到仓日期 / 卸柜日期
    ws.merge_cells("A3:B3")
    ws.cell(row=3, column=1, value="提柜日期")
    ws.merge_cells("C3:D3")
    ws.cell(row=3, column=5, value="到仓日期")
    ws.merge_cells("F3:I3")
    ws.cell(row=3, column=10, value="卸柜日期")
    ws.merge_cells("K3:M3")
    ws.row_dimensions[3].height = 39
    for col in range(1, 14):
        cc = ws.cell(row=3, column=col)
        cc.font = Font(name=FONT_NAME, size=16, bold=False, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.border = thin

    # Row 4: headers
    for col, h in enumerate(BOL_HEADERS, start=1):
        cc = ws.cell(row=4, column=col, value=h)
        cc.font = Font(name=FONT_NAME, size=11, bold=True, color="000000")
        cc.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cc.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cc.border = thin
    ws.row_dimensions[4].height = 48

    # Data rows
    start_row = 5
    for i, r in enumerate(transformed):
        row = start_row + i
        for col_idx, key in enumerate(BOL_KEY_ORDER, start=1):
            ws.cell(row=row, column=col_idx, value=r[key])

        fill_color = _method_fill(r["method"])
        white = _is_white_font(r["method"])
        font = Font(name=FONT_NAME, size=16, bold=True,
                    color=("FFFFFF" if white else "000000"))
        align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        fill = PatternFill("solid", fgColor=fill_color) if fill_color else None
        long = r["address"] and ("\n" in r["address"] or len(r["address"]) > 40)
        for col in range(1, 14):
            cc = ws.cell(row=row, column=col)
            cc.font = font
            cc.alignment = align
            cc.border = thin
            if fill:
                cc.fill = fill
        ws.row_dimensions[row].height = 49 if long else 40

    # Number formats
    for i in range(len(transformed)):
        row = start_row + i
        ws.cell(row=row, column=4).number_format = "0"
        ws.cell(row=row, column=5).number_format = "0.000"
        ws.cell(row=row, column=6).number_format = "0.00"
        ws.cell(row=row, column=7).number_format = "0.000"
        ws.cell(row=row, column=8).number_format = "0.000"

    # Merge FBA CODE
    _merge_consecutive_same(ws, transformed, start_row, col=10, key="fba_code")

    # TOTAL row
    total_row = start_row + len(transformed)
    ws.cell(row=total_row, column=3, value="TOTAL")
    for ci, key in [(4, "ctns"), (5, "cbm"), (6, "kg"), (7, "ctn_lbs"), (8, "total_lbs")]:
        s = round(sum(_to_num(r[key]) for r in transformed), 4)
        ws.cell(row=total_row, column=ci, value=int(s) if s == int(s) else s)

    red = Font(name=FONT_NAME, size=16, bold=True, color="FF0000")
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for col in range(1, 14):
        cc = ws.cell(row=total_row, column=col)
        cc.font = red
        cc.alignment = align
        cc.border = thin
    ws.row_dimensions[total_row].height = 41

    ws.freeze_panes = "A5"
    wb.save(out_path)


def build_bol_packing_list(transformed, container_no, out_path):
    """BOL packing list: same as full list but A-M (13 cols) only."""
    _write_bol_excel(transformed, container_no, container_no, out_path)


def build_bol_hold_list(transformed, container_no, out_path):
    """BOL hold list: only HOLD / RELABEL / PICK UP rows."""
    hold_methods = {"HOLD", "RELABEL", "PICK UP", "PICKUP"}
    filtered = [r for r in transformed if r["method"] in hold_methods]
    if not filtered:
        return None  # No hold data — skip file generation
    _write_bol_excel(filtered, container_no, container_no, out_path)
    return out_path


def build_bol_trucking(transformed, container_no, out_path):
    """BOL trucking: Word doc summarizing pallets and CBM per warehouse/method group."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import math

    doc = Document()
    # Page setup: portrait, letter
    section = doc.sections[0]
    section.page_width = Pt(612)   # 8.5"
    section.page_height = Pt(792)  # 11"

    # Aggregate data by fba_code (warehouse/destination).
    # Every group uses fba_code as the label and shows pallets + CBM.
    # For TRUCK rows, fba_code is the warehouse code (e.g. SBD1, MIT2).
    # For non-TRUCK rows (FEDEX, UPS, LOCAL, FF, etc.), fba_code is the
    # destination name (e.g. "FEDEX", "UPS", "Los Angeles-1", "fulfillment").
    groups = {}   # fba_code -> total_cbm
    group_order = []

    for r in transformed:
        label = r["fba_code"] or r["method"]  # fallback to method if fba_code empty
        cbm = _to_num(r["cbm"])
        if label not in groups:
            groups[label] = 0
            group_order.append(label)
        groups[label] += cbm

    def _add_run(para, text, size_pt, bold=True):
        run = para.add_run(text)
        run.font.name = "SimSun"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Set East Asian font
        from docx.oxml.ns import qn
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), "SimSun")

    # Title
    p = doc.add_paragraph()
    _add_run(p, container_no, 48)

    # Blank line
    doc.add_paragraph()

    # All groups: label-{pallets}P  {cbm}CBM
    for label in group_order:
        cbm = groups[label]
        pallets = max(1, round(cbm / 2))
        cbm_display = f"{cbm:.2f}".rstrip("0").rstrip(".")
        p = doc.add_paragraph()
        _add_run(p, f"{label}-{pallets}P  {cbm_display}CBM", 36)

    doc.save(out_path)


def build_bol_ship_marks(container_no, out_path):
    """BOL 柜号唛头: Word doc with container number and MELKO, landscape."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()
    section = doc.sections[0]
    # Landscape letter
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Pt(792)   # 11" in landscape
    section.page_height = Pt(612)  # 8.5"
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    def _add_run(para, text, size_pt, bold=True):
        run = para.add_run(text)
        run.font.name = "SimSun"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        from docx.oxml.ns import qn
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), "SimSun")

    # Container number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, container_no, 110)

    # Spacer
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, " ", 20, bold=False)

    # MELKO
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p3, "MELKO", 82, bold=False)

    doc.save(out_path)


# =============================================================================
# Orchestrator: build all files at once
# =============================================================================

def build_all(src_path: str, out_folder: str = None):
    """Build all 5 output files from a single source packing list.

    Returns dict of generated file paths.
    """
    # Read and transform
    raw_headers, rows, header_map, header_to_canon, metadata = read_source(src_path)
    container_no = metadata.get("container") or _extract_container(src_path)

    sorted_rows = sort_rows(rows, header_map)
    transformed = [transform_row(r, header_map) for r in sorted_rows]

    # Output folder defaults to same as source
    if out_folder is None:
        out_folder = os.path.dirname(src_path)

    bol_folder = os.path.join(out_folder, "BOL")
    os.makedirs(bol_folder, exist_ok=True)

    # Decide extras for the full packing list
    all_candidates = [
        h for h in raw_headers
        if header_to_canon.get(h) is None
        and h
        and _norm_header_key(h) not in SKIP_HEADER_KEYS
    ]
    extras_headers = [
        h for h in all_candidates
        if any(sr.get(h) not in (None, "") for sr in sorted_rows)
    ]

    # 1. Full packing list (19-col)
    full_path = os.path.join(out_folder, f"{container_no} packing list.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "总表"
    for col_letter, w in COLUMN_WIDTHS.items():
        ws.column_dimensions[col_letter].width = w
    _write_header_block(ws, container_no, metadata, extras_headers)
    last_data = _write_data_rows(ws, transformed, sorted_rows, extras_headers)
    _write_total_row(ws, last_data, transformed, len(extras_headers))
    ws.freeze_panes = "A4"
    wb.save(full_path)

    # 2. BOL packing list (A-M only)
    bol_pl = os.path.join(bol_folder, f"{container_no} packing list.xlsx")
    build_bol_packing_list(transformed, container_no, bol_pl)

    # 3. BOL hold list (only if there are HOLD/RELABEL/PICK UP rows)
    bol_hl = os.path.join(bol_folder, f"{container_no} hold list.xlsx")
    bol_hl_result = build_bol_hold_list(transformed, container_no, bol_hl)
    if bol_hl_result is None:
        bol_hl = None  # not generated

    # 4. BOL trucking
    bol_tr = os.path.join(bol_folder, f"{container_no}-TRUCKING.docx")
    build_bol_trucking(transformed, container_no, bol_tr)

    # 5. BOL ship marks
    bol_sm = os.path.join(bol_folder, f"{container_no}-柜号唛头.docx")
    build_bol_ship_marks(container_no, bol_sm)

    return {
        "full_packing_list": full_path,
        "bol_packing_list": bol_pl,
        "bol_hold_list": bol_hl,
        "bol_trucking": bol_tr,
        "bol_ship_marks": bol_sm,
    }


# =============================================================================
# CLI
# =============================================================================

def _default_output_path(src_path: str) -> str:
    container = _extract_container(src_path)
    folder = os.path.dirname(src_path)
    return os.path.join(folder, f"{container} packing list.xlsx")


def main():
    import sys
    if len(sys.argv) < 2:
        print("Usage: python converter.py <source.xlsx|.xls> [output_folder]")
        sys.exit(1)
    src = sys.argv[1]
    out_folder = sys.argv[2] if len(sys.argv) >= 3 else None
    results = build_all(src, out_folder)
    for label, path in results.items():
        print(f"[OK] {label}: {os.path.basename(path)}")


if __name__ == "__main__":
    main()
